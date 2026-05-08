from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, List, Optional
from uuid import uuid4

from fastapi import HTTPException

MAX_DOCUMENT_BYTES = 20 * 1024 * 1024
INDEXABLE_EXTENSIONS = {".pdf", ".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

# Minimum character threshold to consider extraction successful
MIN_EXTRACTION_CHARS = 50


@dataclass
class IndexedDocumentPayload:
    original_name: str
    extension: str
    stored_name: str
    file_url: str
    pages: List[dict]
    documents: List[str]
    metadatas: List[dict]
    extracted_text: str = ""
    extraction_warnings: List[str] = field(default_factory=list)

    @property
    def source_type(self) -> str:
        return "pdf" if self.extension == ".pdf" else "word"

    @property
    def is_extraction_weak(self) -> bool:
        return len(self.extracted_text.strip()) < MIN_EXTRACTION_CHARS


def _clean_arabic_text(text: str) -> str:
    """Clean Arabic text by removing RTL/LTR markers and normalizing whitespace."""
    if not text:
        return ""
    # Remove Unicode directional markers
    cleaned = text.replace("\u200f", " ").replace("\u200e", " ")
    # Remove zero-width chars
    cleaned = cleaned.replace("\u200b", "").replace("\u200c", "").replace("\u200d", "")
    cleaned = cleaned.replace("\ufeff", "")
    # Normalize whitespace
    cleaned = " ".join(cleaned.split())
    return cleaned.strip()


def ensure_supported_upload(filename: str) -> str:
    original_name = Path(str(filename or "").strip()).name
    if not original_name:
        raise HTTPException(status_code=400, detail="File name is required")

    ext = Path(original_name).suffix.lower()
    if ext == ".doc":
        raise HTTPException(status_code=400, detail="DOC files are not supported. Please upload DOCX or PDF.")
    if ext not in {*INDEXABLE_EXTENSIONS, *IMAGE_EXTENSIONS}:
        raise HTTPException(
            status_code=400,
            detail="Only PDF, DOCX, PNG, JPG, JPEG, WEBP, or GIF files are supported",
        )
    return ext


def ensure_upload_content(content: bytes) -> None:
    if not content:
        raise HTTPException(status_code=400, detail="Empty file is not allowed")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise HTTPException(status_code=413, detail="File is too large (max 20MB)")


def save_uploaded_binary(content: bytes, original_name: str, storage_dir: Path) -> tuple[str, Path]:
    ext = Path(original_name).suffix.lower() or ".bin"
    stored_name = f"{uuid4().hex}{ext}"
    destination = storage_dir / stored_name
    destination.write_bytes(content)
    return stored_name, destination


def extract_document_pages(content: bytes, extension: str) -> List[dict]:
    if extension == ".pdf":
        return _pdf_pages_from_upload(content)
    if extension == ".docx":
        return _docx_pages_from_upload(content)
    raise HTTPException(status_code=400, detail="Only PDF or DOCX files can be indexed.")


def chunk_document_pages(pages: List[dict], chunk_size: int = 900, chunk_overlap: int = 120) -> tuple[List[str], List[dict]]:
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
    except Exception:
        try:
            from langchain.text_splitter import RecursiveCharacterTextSplitter
        except Exception as exc:
            raise HTTPException(status_code=500, detail="Text splitter dependency is missing on the server.") from exc

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    documents: List[str] = []
    metadatas: List[dict] = []
    for item in pages:
        chunks = splitter.split_text(item["text"])
        for chunk_idx, chunk in enumerate(chunks, start=1):
            clean = _clean_arabic_text(chunk)
            if not clean:
                continue
            documents.append(clean)
            metadatas.append({"page": item["page"], "chunk": chunk_idx})
    return documents, metadatas


def _build_full_extracted_text(pages: List[dict]) -> str:
    """Build full extracted text from all pages for preview storage."""
    parts = []
    for page in pages:
        text = _clean_arabic_text(page.get("text", ""))
        if text:
            parts.append(f"--- صفحة {page.get('page', '?')} ---\n{text}")
    return "\n\n".join(parts)


def _check_extraction_quality(extracted_text: str, original_name: str) -> List[str]:
    """Check extraction quality and return warnings."""
    warnings = []
    text = extracted_text.strip()

    if not text:
        warnings.append(f"⚠️ لم يتم استخراج أي نص من الملف '{original_name}'. الملف قد يكون صورة فقط (scanned PDF).")
        return warnings

    if len(text) < MIN_EXTRACTION_CHARS:
        warnings.append(
            f"⚠️ النص المستخرج قصير جدًا ({len(text)} حرف فقط). "
            "قد يكون الملف يحتوي على صور أو نص غير قابل للاستخراج."
        )

    # Check for Arabic content
    arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F')
    total_chars = len(text)
    if total_chars > 0 and arabic_chars / total_chars < 0.1:
        # Low Arabic ratio might indicate encoding issues
        latin_chars = sum(1 for c in text if c.isascii() and c.isalpha())
        if latin_chars / total_chars > 0.5:
            pass  # Probably English document, that's fine
        else:
            warnings.append(
                "⚠️ نسبة النص العربي منخفضة. قد يكون هناك مشكلة في استخراج النص. "
                "يُفضل استخدام DOCX بدلاً من PDF للنصوص العربية."
            )

    # Check for garbled text (lots of replacement characters)
    replacement_chars = text.count('\ufffd') + text.count('?') + text.count('□')
    if total_chars > 0 and replacement_chars / total_chars > 0.05:
        warnings.append(
            "⚠️ يوجد نسبة عالية من الأحرف غير المقروءة. "
            "يُفضل رفع الملف بصيغة DOCX أو إدخال النص يدويًا عبر knowledge_text."
        )

    return warnings


def _docx_namespace() -> dict[str, str]:
    return {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def _docx_local_name(tag: str) -> str:
    return str(tag or "").split("}", 1)[-1]


def _docx_join_text_nodes(element: ET.Element, namespace: dict[str, str]) -> str:
    fragments: List[str] = []
    for node in element.iter():
        local_name = _docx_local_name(node.tag)
        if local_name == "t" and node.text:
            fragments.append(node.text)
        elif local_name == "tab":
            fragments.append("\t")
        elif local_name in {"br", "cr"}:
            fragments.append("\n")
    return _clean_arabic_text("".join(fragments))


def _docx_extract_table_lines(table: ET.Element, namespace: dict[str, str]) -> List[str]:
    lines: List[str] = []
    for row in table.findall("./w:tr", namespace):
        cells: List[str] = []
        for cell in row.findall("./w:tc", namespace):
            fragments: List[str] = []
            for paragraph in cell.findall(".//w:p", namespace):
                text = _docx_join_text_nodes(paragraph, namespace)
                if text:
                    fragments.append(text)
            cell_text = _clean_arabic_text(" ".join(fragments))
            if cell_text:
                cells.append(cell_text)
        if cells:
            lines.append(" | ".join(cells))
    return lines


def _docx_extract_block_lines(root: ET.Element) -> List[str]:
    namespace = _docx_namespace()
    container = None
    for path in ("./w:body", "./w:hdr", "./w:ftr", "./w:footnotes", "./w:endnotes"):
        container = root.find(path, namespace)
        if container is not None:
            break
    if container is None:
        container = root

    lines: List[str] = []
    for child in list(container):
        local_name = _docx_local_name(child.tag)
        if local_name == "p":
            text = _docx_join_text_nodes(child, namespace)
            if text:
                lines.append(text)
        elif local_name == "tbl":
            lines.extend(_docx_extract_table_lines(child, namespace))

    if lines:
        return lines

    fallback_lines: List[str] = []
    for paragraph in root.findall(".//w:p", namespace):
        text = _docx_join_text_nodes(paragraph, namespace)
        if text:
            fallback_lines.append(text)
    return fallback_lines


def _docx_relevant_part_names(archive: zipfile.ZipFile) -> List[str]:
    preferred = [
        "word/document.xml",
        "word/footnotes.xml",
        "word/endnotes.xml",
    ]
    names = set(archive.namelist())
    ordered = [name for name in preferred if name in names]
    ordered.extend(sorted(name for name in names if name.startswith("word/header") and name.endswith(".xml")))
    ordered.extend(sorted(name for name in names if name.startswith("word/footer") and name.endswith(".xml")))
    return ordered


def _merge_unique_docx_lines(parts: Iterable[str]) -> List[str]:
    merged: List[str] = []
    seen: set[str] = set()
    for part in parts:
        clean = _clean_arabic_text(part)
        if not clean:
            continue
        if clean in seen:
            continue
        merged.append(clean)
        seen.add(clean)
    return merged


def _extract_docx_text_lines(content: bytes) -> List[str]:
    primary_lines: List[str] = []

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(BytesIO(content))
        for para in doc.paragraphs:
            text = _clean_arabic_text(para.text)
            if text:
                primary_lines.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [_clean_arabic_text(cell.text) for cell in row.cells if _clean_arabic_text(cell.text)]
                if cells:
                    primary_lines.append(" | ".join(cells))

        for section in doc.sections:
            for container in (section.header, section.footer):
                for para in getattr(container, "paragraphs", []):
                    text = _clean_arabic_text(getattr(para, "text", ""))
                    if text:
                        primary_lines.append(text)
                for table in getattr(container, "tables", []):
                    for row in table.rows:
                        cells = [_clean_arabic_text(cell.text) for cell in row.cells if _clean_arabic_text(cell.text)]
                        if cells:
                            primary_lines.append(" | ".join(cells))
    except ImportError:
        pass

    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            part_names = _docx_relevant_part_names(archive)
            if not part_names:
                raise HTTPException(status_code=400, detail="The DOCX file does not contain readable XML parts.")

            supplemental_lines: List[str] = []
            for part_name in part_names:
                try:
                    root = ET.fromstring(archive.read(part_name))
                except ET.ParseError as exc:
                    raise HTTPException(status_code=400, detail=f"Could not parse DOCX part: {part_name}") from exc
                supplemental_lines.extend(_docx_extract_block_lines(root))

            merged_lines = _merge_unique_docx_lines(primary_lines)
            seen = set(merged_lines)
            for line in supplemental_lines:
                clean = _clean_arabic_text(line)
                if not clean or clean in seen:
                    continue
                merged_lines.append(clean)
                seen.add(clean)
            return merged_lines
    except KeyError as exc:
        raise HTTPException(status_code=400, detail="The DOCX file does not contain a readable document body.") from exc
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=400, detail="Invalid DOCX file.") from exc


def prepare_indexable_document(content: bytes, original_name: str, storage_dir: Path) -> IndexedDocumentPayload:
    extension = Path(original_name).suffix.lower()
    stored_name, _ = save_uploaded_binary(content, original_name, storage_dir)
    pages = extract_document_pages(content, extension)

    # Build full extracted text even if pages are empty (for diagnostics)
    extracted_text = _build_full_extracted_text(pages) if pages else ""
    warnings = _check_extraction_quality(extracted_text, original_name)

    if not pages:
        raise HTTPException(
            status_code=400,
            detail="No extractable text was found in this file. "
                   "If this is a scanned PDF, try uploading a DOCX version or entering text manually."
        )

    documents, metadatas = chunk_document_pages(pages)
    if not documents:
        raise HTTPException(
            status_code=400,
            detail="No extractable text chunks were found in this file."
        )

    return IndexedDocumentPayload(
        original_name=original_name,
        extension=extension,
        stored_name=stored_name,
        file_url=f"/api/storage/files/{stored_name}",
        pages=pages,
        documents=documents,
        metadatas=metadatas,
        extracted_text=extracted_text,
        extraction_warnings=warnings,
    )


def prepare_indexable_document_from_existing(content: bytes, original_name: str, stored_name: str) -> IndexedDocumentPayload:
    extension = Path(original_name).suffix.lower()
    pages = extract_document_pages(content, extension)

    extracted_text = _build_full_extracted_text(pages) if pages else ""
    warnings = _check_extraction_quality(extracted_text, original_name)

    if not pages:
        raise HTTPException(
            status_code=400,
            detail="No extractable text was found in this file. "
                   "If this is a scanned PDF, try uploading a DOCX version or entering text manually."
        )

    documents, metadatas = chunk_document_pages(pages)
    if not documents:
        raise HTTPException(status_code=400, detail="No extractable text chunks were found in this file.")

    return IndexedDocumentPayload(
        original_name=original_name,
        extension=extension,
        stored_name=stored_name,
        file_url=f"/api/storage/files/{stored_name}",
        pages=pages,
        documents=documents,
        metadatas=metadatas,
        extracted_text=extracted_text,
        extraction_warnings=warnings,
    )


def index_prepared_document(rag_chatbot: Any, prepared: IndexedDocumentPayload, base_metadata: dict) -> None:
    if rag_chatbot is None:
        raise HTTPException(status_code=503, detail="AI Service unavailable.")

    rag_chatbot.index_documents(
        prepared.documents,
        metadatas=[{**base_metadata, **meta} for meta in prepared.metadatas],
    )
    rag_chatbot.flush()


def delete_document_vectors(rag_chatbot: Any, document_id: str) -> int:
    """Delete all vectors for a given document_id. Returns count of deleted vectors."""
    if rag_chatbot is None or not document_id:
        return 0

    vector_store = getattr(rag_chatbot, "vector_store", None)
    if vector_store is None:
        return 0

    deleted = 0
    try:
        col = getattr(vector_store, "_collection", None)
        if col is not None:
            existing = col.get(where={"document_id": str(document_id)})
            old_count = len(existing.get("ids", []))
            if old_count > 0:
                col.delete(ids=existing["ids"])
                deleted = old_count
        elif hasattr(vector_store, "delete"):
            vector_store.delete(where={"document_id": str(document_id)})
            deleted = -1  # Unknown count
    except Exception:
        pass

    return deleted


def _pdf_pages_from_upload(content: bytes) -> List[dict]:
    pages: List[dict] = []
    
    # Try PyMuPDF (fitz) first - much better for Arabic and complex layouts
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        for idx, page in enumerate(doc, start=1):
            text = (page.get_text() or "").strip()
            if text:
                text = _clean_arabic_text(text)
                if text:
                    pages.append({"page": idx, "text": text})
        if pages:
            return pages
    except Exception:
        pass  # Fallback to pypdf if fitz fails or returns empty

    # Fallback to pypdf
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise HTTPException(status_code=500, detail="PDF libraries (fitz, pypdf) are not installed on the server.") from exc

    try:
        reader = PdfReader(BytesIO(content))
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            text = _clean_arabic_text(text)
            if text:
                pages.append({"page": idx, "text": text})
    except Exception:
        pass

    return pages


def _docx_pages_from_upload(content: bytes) -> List[dict]:
    lines = _extract_docx_text_lines(content)
    merged = "\n".join(lines).strip()
    if not merged:
        return []
    return [{"page": 1, "text": merged}]
